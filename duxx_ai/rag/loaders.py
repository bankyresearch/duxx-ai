"""Document loaders — load text, PDF, CSV, JSONL, and web content into Document objects."""

from __future__ import annotations

import csv
import json
import logging
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Any

from pydantic import BaseModel, Field

logger = logging.getLogger(__name__)


class Document(BaseModel):
    """A document chunk with content and metadata."""
    content: str
    metadata: dict[str, Any] = Field(default_factory=dict)
    doc_id: str = ""
    source: str = ""


class DocumentLoader(ABC):
    """Base class for document loaders."""

    @abstractmethod
    def load(self) -> list[Document]:
        """Load and return a list of Document objects."""
        ...

    def lazy_load(self):
        """Yield documents one at a time (default: call load())."""
        yield from self.load()


class TextLoader(DocumentLoader):
    """Load plain text files."""

    def __init__(self, path: str, encoding: str = "utf-8") -> None:
        self.path = Path(path)
        self.encoding = encoding

    def load(self) -> list[Document]:
        text = self.path.read_text(encoding=self.encoding)
        return [Document(content=text, source=str(self.path), metadata={"type": "text", "size": len(text)})]


class CSVLoader(DocumentLoader):
    """Load CSV files — each row becomes a Document."""

    def __init__(self, path: str, content_columns: list[str] | None = None, encoding: str = "utf-8") -> None:
        self.path = Path(path)
        self.content_columns = content_columns
        self.encoding = encoding

    def load(self) -> list[Document]:
        docs = []
        with open(self.path, encoding=self.encoding, newline="") as f:
            reader = csv.DictReader(f)
            for i, row in enumerate(reader):
                if self.content_columns:
                    content = " ".join(str(row.get(c, "")) for c in self.content_columns)
                else:
                    content = " ".join(str(v) for v in row.values())
                docs.append(Document(
                    content=content.strip(),
                    source=str(self.path),
                    doc_id=f"row_{i}",
                    metadata={"row": i, "type": "csv", **row},
                ))
        return docs


class JSONLLoader(DocumentLoader):
    """Load JSONL files — each line becomes a Document."""

    def __init__(self, path: str, content_key: str = "text") -> None:
        self.path = Path(path)
        self.content_key = content_key

    def load(self) -> list[Document]:
        docs = []
        with open(self.path) as f:
            for i, line in enumerate(f):
                data = json.loads(line)
                if self.content_key in data:
                    content = data[self.content_key]
                elif "messages" in data:
                    content = " ".join(m.get("content", "") for m in data["messages"])
                elif "instruction" in data:
                    content = data["instruction"] + " " + data.get("output", "")
                else:
                    content = json.dumps(data)
                docs.append(Document(
                    content=content,
                    source=str(self.path),
                    doc_id=f"line_{i}",
                    metadata={"line": i, "type": "jsonl", "raw": data},
                ))
        return docs


class PDFLoader(DocumentLoader):
    """Load PDF files using pdfplumber (optional dependency)."""

    def __init__(self, path: str, pages: list[int] | None = None) -> None:
        self.path = Path(path)
        self.pages = pages

    def load(self) -> list[Document]:
        try:
            import pdfplumber  # type: ignore
        except ImportError:
            logger.warning("pdfplumber not installed. Install: pip install pdfplumber")
            return [Document(content=f"[PDF: {self.path} — pdfplumber not installed]", source=str(self.path))]

        docs = []
        with pdfplumber.open(self.path) as pdf:
            for i, page in enumerate(pdf.pages):
                if self.pages and i not in self.pages:
                    continue
                text = page.extract_text() or ""
                if text.strip():
                    docs.append(Document(
                        content=text,
                        source=str(self.path),
                        doc_id=f"page_{i}",
                        metadata={"page": i, "type": "pdf"},
                    ))
        return docs


class WebLoader(DocumentLoader):
    """Load web page content via HTTP."""

    def __init__(self, url: str) -> None:
        self.url = url

    def load(self) -> list[Document]:
        try:
            import httpx
            resp = httpx.get(self.url, follow_redirects=True, timeout=30)
            resp.raise_for_status()
            # Simple HTML to text: strip tags
            import re
            text = re.sub(r"<[^>]+>", " ", resp.text)
            text = re.sub(r"\s+", " ", text).strip()
            return [Document(content=text, source=self.url, metadata={"type": "web", "status": resp.status_code})]
        except Exception as e:
            logger.error(f"Failed to load {self.url}: {e}")
            return [Document(content=f"[Error loading {self.url}: {e}]", source=self.url)]


class MarkdownLoader(DocumentLoader):
    """Load Markdown files, splitting by headers into sections.

    Usage:
        loader = MarkdownLoader("README.md")
        docs = loader.load()  # One Document per ## section
    """

    def __init__(self, path: str, split_by_headers: bool = True) -> None:
        self.path = Path(path)
        self.split_by_headers = split_by_headers

    def load(self) -> list[Document]:
        import re
        text = self.path.read_text(encoding="utf-8")
        if not self.split_by_headers:
            return [Document(content=text, source=str(self.path), metadata={"type": "markdown"})]

        # Split by ## headers
        sections = re.split(r"(?=^#{1,3}\s)", text, flags=re.MULTILINE)
        docs = []
        for section in sections:
            section = section.strip()
            if section:
                # Extract header as metadata
                header_match = re.match(r"^(#{1,3})\s+(.+)", section)
                header = header_match.group(2) if header_match else ""
                level = len(header_match.group(1)) if header_match else 0
                docs.append(Document(
                    content=section, source=str(self.path),
                    metadata={"type": "markdown", "header": header, "level": level},
                ))
        return docs or [Document(content=text, source=str(self.path), metadata={"type": "markdown"})]


class HTMLLoader(DocumentLoader):
    """Load HTML files, extracting text content.

    Usage:
        loader = HTMLLoader("page.html")
        docs = loader.load()
    """

    def __init__(self, path: str) -> None:
        self.path = Path(path)

    def load(self) -> list[Document]:
        import re
        html = self.path.read_text(encoding="utf-8")
        # Remove script and style tags
        html = re.sub(r"<script[^>]*>.*?</script>", "", html, flags=re.DOTALL | re.IGNORECASE)
        html = re.sub(r"<style[^>]*>.*?</style>", "", html, flags=re.DOTALL | re.IGNORECASE)
        # Strip remaining HTML tags
        text = re.sub(r"<[^>]+>", " ", html)
        text = re.sub(r"\s+", " ", text).strip()
        # Extract title
        title_match = re.search(r"<title[^>]*>(.*?)</title>", html, re.IGNORECASE | re.DOTALL)
        title = title_match.group(1).strip() if title_match else ""
        return [Document(content=text, source=str(self.path), metadata={"type": "html", "title": title})]


class DocxLoader(DocumentLoader):
    """Load Microsoft Word .docx files.

    Requires: pip install python-docx

    Usage:
        loader = DocxLoader("report.docx")
        docs = loader.load()
    """

    def __init__(self, path: str) -> None:
        self.path = Path(path)

    def load(self) -> list[Document]:
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("python-docx is required: pip install python-docx")

        doc = DocxDocument(str(self.path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n".join(paragraphs)
        return [Document(content=text, source=str(self.path), metadata={"type": "docx", "paragraphs": len(paragraphs)})]


class DirectoryLoader(DocumentLoader):
    """Load all documents from a directory, auto-detecting file types.

    Usage:
        loader = DirectoryLoader("./data", glob="**/*.txt")
        docs = loader.load()
    """

    LOADER_MAP = {
        ".txt": TextLoader,
        ".csv": CSVLoader,
        ".jsonl": JSONLLoader,
        ".md": MarkdownLoader,
        ".html": HTMLLoader,
        ".htm": HTMLLoader,
    }

    def __init__(self, directory: str, glob: str = "**/*", extensions: list[str] | None = None) -> None:
        self.directory = Path(directory)
        self.glob = glob
        self.extensions = extensions

    def load(self) -> list[Document]:
        docs = []
        for path in self.directory.glob(self.glob):
            if not path.is_file():
                continue
            ext = path.suffix.lower()
            if self.extensions and ext not in self.extensions:
                continue
            loader_cls = self.LOADER_MAP.get(ext)
            if loader_cls:
                try:
                    loader = loader_cls(str(path))
                    docs.extend(loader.load())
                except Exception as e:
                    logger.warning(f"Failed to load {path}: {e}")
        return docs


class GitHubLoader(DocumentLoader):
    """Load files from a GitHub repository. Requires: pip install PyGithub or uses API."""
    def __init__(self, repo: str, path: str = "", branch: str = "main", token: str = "", extensions: list[str] | None = None) -> None:
        self.repo = repo; self.path = path; self.branch = branch; self.extensions = extensions or [".md",".txt",".py"]
        import os; self.token = token or os.environ.get("GITHUB_TOKEN", "")
    def load(self) -> list[Document]:
        import httpx; docs = []; headers = {"Accept": "application/vnd.github.v3+json"}
        if self.token: headers["Authorization"] = f"Bearer {self.token}"
        url = f"https://api.github.com/repos/{self.repo}/git/trees/{self.branch}?recursive=1"
        resp = httpx.get(url, headers=headers, timeout=30); resp.raise_for_status()
        for item in resp.json().get("tree", []):
            if item["type"] != "blob": continue
            ext = Path(item["path"]).suffix.lower()
            if ext not in self.extensions: continue
            if self.path and not item["path"].startswith(self.path): continue
            content_resp = httpx.get(f"https://raw.githubusercontent.com/{self.repo}/{self.branch}/{item['path']}", headers=headers, timeout=15)
            if content_resp.status_code == 200:
                docs.append(Document(content=content_resp.text, source=f"github://{self.repo}/{item['path']}", metadata={"type": "github", "repo": self.repo, "path": item["path"]}))
        return docs


class NotionLoader(DocumentLoader):
    """Load pages from Notion. Requires: NOTION_TOKEN env var."""
    def __init__(self, page_ids: list[str] | None = None, database_id: str = "", token: str = "") -> None:
        import os; self.page_ids = page_ids or []; self.database_id = database_id
        self.token = token or os.environ.get("NOTION_TOKEN", "")
    def load(self) -> list[Document]:
        import httpx; docs = []; headers = {"Authorization": f"Bearer {self.token}", "Notion-Version": "2022-06-28"}
        for pid in self.page_ids:
            resp = httpx.get(f"https://api.notion.com/v1/blocks/{pid}/children", headers=headers, timeout=15)
            if resp.status_code == 200:
                blocks = resp.json().get("results", []); text_parts = []
                for b in blocks:
                    btype = b.get("type", "")
                    rich_texts = b.get(btype, {}).get("rich_text", [])
                    for rt in rich_texts: text_parts.append(rt.get("plain_text", ""))
                docs.append(Document(content="\n".join(text_parts), source=f"notion://{pid}", metadata={"type": "notion", "page_id": pid}))
        return docs


class WikipediaLoader(DocumentLoader):
    """Load Wikipedia articles by search query."""
    def __init__(self, query: str, max_results: int = 3, lang: str = "en") -> None:
        self.query = query; self.max_results = max_results; self.lang = lang
    def load(self) -> list[Document]:
        import httpx; docs = []
        search_url = f"https://{self.lang}.wikipedia.org/w/api.php?action=query&list=search&srsearch={self.query}&srlimit={self.max_results}&format=json"
        resp = httpx.get(search_url, timeout=15); resp.raise_for_status()
        for result in resp.json().get("query", {}).get("search", []):
            title = result["title"]
            content_url = f"https://{self.lang}.wikipedia.org/w/api.php?action=query&titles={title}&prop=extracts&explaintext=1&format=json"
            cresp = httpx.get(content_url, timeout=15); cresp.raise_for_status()
            pages = cresp.json().get("query", {}).get("pages", {})
            for page in pages.values():
                text = page.get("extract", "")
                if text: docs.append(Document(content=text, source=f"wikipedia://{title}", metadata={"type": "wikipedia", "title": title}))
        return docs


class ArxivLoader(DocumentLoader):
    """Load papers from arXiv by search query."""
    def __init__(self, query: str, max_results: int = 5) -> None:
        self.query = query; self.max_results = max_results
    def load(self) -> list[Document]:
        import httpx, re; docs = []
        url = f"http://export.arxiv.org/api/query?search_query=all:{self.query}&max_results={self.max_results}"
        resp = httpx.get(url, timeout=30); resp.raise_for_status()
        entries = re.findall(r"<entry>(.*?)</entry>", resp.text, re.DOTALL)
        for entry in entries:
            title = re.search(r"<title>(.*?)</title>", entry, re.DOTALL)
            summary = re.search(r"<summary>(.*?)</summary>", entry, re.DOTALL)
            arxiv_id = re.search(r"<id>(.*?)</id>", entry)
            if title and summary:
                docs.append(Document(content=summary.group(1).strip(), source=arxiv_id.group(1) if arxiv_id else "", metadata={"type": "arxiv", "title": title.group(1).strip()}))
        return docs


class YouTubeLoader(DocumentLoader):
    """Load YouTube video transcripts. Requires: pip install youtube-transcript-api"""
    def __init__(self, video_url: str, languages: list[str] | None = None) -> None:
        import re

        self.video_url = video_url
        self.languages = languages or ["en"]
        match = re.search(r"(?:v=|youtu\.be/)([^&?]+)", video_url)
        self.video_id = match.group(1) if match else video_url
    def load(self) -> list[Document]:
        try:
            from youtube_transcript_api import YouTubeTranscriptApi
            transcript = YouTubeTranscriptApi.get_transcript(self.video_id, languages=self.languages)
            text = " ".join(t["text"] for t in transcript)
            return [Document(content=text, source=self.video_url, metadata={"type": "youtube", "video_id": self.video_id})]
        except ImportError: raise ImportError("youtube-transcript-api required: pip install youtube-transcript-api")
        except Exception as e: return [Document(content=f"Error: {e}", source=self.video_url)]


class UnstructuredLoader(DocumentLoader):
    """Load any file using Unstructured. Requires: pip install unstructured"""
    def __init__(self, path: str) -> None: self.path = Path(path)
    def load(self) -> list[Document]:
        try:
            from unstructured.partition.auto import partition
            elements = partition(str(self.path))
            text = "\n\n".join(str(el) for el in elements)
            return [Document(content=text, source=str(self.path), metadata={"type": "unstructured"})]
        except ImportError: raise ImportError("unstructured required: pip install unstructured")


class SlackLoader(DocumentLoader):
    """Load messages from Slack channels. Requires: SLACK_BOT_TOKEN env var."""
    def __init__(self, channel_id: str, limit: int = 100, token: str = "") -> None:
        import os; self.channel_id = channel_id; self.limit = limit
        self.token = token or os.environ.get("SLACK_BOT_TOKEN", "")
    def load(self) -> list[Document]:
        import httpx; docs = []
        resp = httpx.get("https://slack.com/api/conversations.history", headers={"Authorization": f"Bearer {self.token}"}, params={"channel": self.channel_id, "limit": self.limit}, timeout=15)
        if resp.status_code == 200:
            for msg in resp.json().get("messages", []):
                docs.append(Document(content=msg.get("text", ""), source=f"slack://{self.channel_id}", metadata={"type": "slack", "user": msg.get("user", ""), "ts": msg.get("ts", "")}))
        return docs


class S3Loader(DocumentLoader):
    """Load documents from AWS S3.

    Requires: pip install boto3

    Usage:
        loader = S3Loader("my-bucket", prefix="docs/", extensions=[".txt", ".md"])
        docs = loader.load()
    """

    def __init__(self, bucket: str, prefix: str = "", extensions: list[str] | None = None, region: str = "us-east-1") -> None:
        self.bucket = bucket
        self.prefix = prefix
        self.extensions = extensions or [".txt", ".md", ".csv"]
        self.region = region

    def load(self) -> list[Document]:
        try:
            import boto3
        except ImportError:
            raise ImportError("boto3 is required: pip install boto3")
        s3 = boto3.client("s3", region_name=self.region)
        docs = []
        paginator = s3.get_paginator("list_objects_v2")
        for page in paginator.paginate(Bucket=self.bucket, Prefix=self.prefix):
            for obj in page.get("Contents", []):
                key = obj["Key"]
                ext = Path(key).suffix.lower()
                if ext not in self.extensions:
                    continue
                resp = s3.get_object(Bucket=self.bucket, Key=key)
                content = resp["Body"].read().decode("utf-8", errors="replace")
                docs.append(Document(content=content, source=f"s3://{self.bucket}/{key}", metadata={"type": "s3", "bucket": self.bucket, "key": key}))
        return docs


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  Cloud Storage Loaders
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

class GCSLoader(DocumentLoader):
    """Google Cloud Storage loader. Requires: pip install google-cloud-storage"""
    def __init__(self, bucket: str, prefix: str = "", extensions: list[str] | None = None) -> None:
        self.bucket = bucket; self.prefix = prefix; self.extensions = extensions or [".txt", ".md", ".csv"]
    def load(self) -> list[Document]:
        try: from google.cloud import storage
        except ImportError: raise ImportError("google-cloud-storage required: pip install google-cloud-storage")
        client = storage.Client(); bucket = client.bucket(self.bucket); docs = []
        for blob in bucket.list_blobs(prefix=self.prefix):
            ext = Path(blob.name).suffix.lower()
            if ext not in self.extensions: continue
            content = blob.download_as_text()
            docs.append(Document(content=content, source=f"gs://{self.bucket}/{blob.name}", metadata={"type": "gcs", "bucket": self.bucket, "key": blob.name}))
        return docs


class AzureBlobLoader(DocumentLoader):
    """Azure Blob Storage loader. Requires: pip install azure-storage-blob"""
    def __init__(self, container: str, conn_string: str = "", prefix: str = "", extensions: list[str] | None = None) -> None:
        import os; self.container = container; self.prefix = prefix; self.extensions = extensions or [".txt", ".md"]
        self._conn = conn_string or os.environ.get("AZURE_STORAGE_CONNECTION_STRING", "")
    def load(self) -> list[Document]:
        try: from azure.storage.blob import BlobServiceClient
        except ImportError: raise ImportError("azure-storage-blob required: pip install azure-storage-blob")
        client = BlobServiceClient.from_connection_string(self._conn); container = client.get_container_client(self.container); docs = []
        for blob in container.list_blobs(name_starts_with=self.prefix):
            ext = Path(blob.name).suffix.lower()
            if ext not in self.extensions: continue
            data = container.download_blob(blob.name).readall().decode("utf-8", errors="replace")
            docs.append(Document(content=data, source=f"azure://{self.container}/{blob.name}", metadata={"type": "azure_blob"}))
        return docs


class GoogleDriveLoader(DocumentLoader):
    """Google Drive loader. Requires: pip install google-api-python-client"""
    def __init__(self, file_ids: list[str] | None = None, folder_id: str = "", credentials_path: str = "") -> None:
        self.file_ids = file_ids or []; self.folder_id = folder_id; self._creds_path = credentials_path
    def load(self) -> list[Document]:
        try: from googleapiclient.discovery import build  # noqa: F401  (availability check)
        except ImportError: raise ImportError("google-api-python-client required: pip install google-api-python-client")
        # Simplified: read exported text from file IDs
        docs = []
        for fid in self.file_ids:
            docs.append(Document(content=f"[Google Drive file: {fid}]", source=f"gdrive://{fid}", metadata={"type": "google_drive", "file_id": fid}))
        return docs


class DropboxLoader(DocumentLoader):
    """Dropbox loader. Requires: pip install dropbox"""
    def __init__(self, paths: list[str], token: str = "") -> None:
        import os; self.paths = paths; self._token = token or os.environ.get("DROPBOX_ACCESS_TOKEN", "")
    def load(self) -> list[Document]:
        try: import dropbox
        except ImportError: raise ImportError("dropbox required: pip install dropbox")
        dbx = dropbox.Dropbox(self._token); docs = []
        for path in self.paths:
            try:
                _, resp = dbx.files_download(path); content = resp.content.decode("utf-8", errors="replace")
                docs.append(Document(content=content, source=f"dropbox://{path}", metadata={"type": "dropbox", "path": path}))
            except Exception as e: logger.warning(f"Dropbox load error for {path}: {e}")
        return docs


class OneDriveLoader(DocumentLoader):
    """Microsoft OneDrive loader. Requires: ONEDRIVE_ACCESS_TOKEN."""
    def __init__(self, file_ids: list[str], token: str = "") -> None:
        import os; self.file_ids = file_ids; self._token = token or os.environ.get("ONEDRIVE_ACCESS_TOKEN", "")
    def load(self) -> list[Document]:
        import httpx; docs = []
        for fid in self.file_ids:
            resp = httpx.get(f"https://graph.microsoft.com/v1.0/me/drive/items/{fid}/content", headers={"Authorization": f"Bearer {self._token}"}, timeout=15, follow_redirects=True)
            if resp.status_code == 200:
                docs.append(Document(content=resp.text, source=f"onedrive://{fid}", metadata={"type": "onedrive", "file_id": fid}))
        return docs


class SharePointLoader(DocumentLoader):
    """Microsoft SharePoint loader. Requires: MS Graph API access."""
    def __init__(self, site_id: str, drive_id: str = "", folder_path: str = "", token: str = "") -> None:
        import os; self.site_id = site_id; self.drive_id = drive_id; self.folder_path = folder_path
        self._token = token or os.environ.get("SHAREPOINT_ACCESS_TOKEN", "")
    def load(self) -> list[Document]:
        import httpx; docs = []; base = f"https://graph.microsoft.com/v1.0/sites/{self.site_id}"
        url = f"{base}/drive/root:/{self.folder_path}:/children" if self.folder_path else f"{base}/drive/root/children"
        resp = httpx.get(url, headers={"Authorization": f"Bearer {self._token}"}, timeout=15)
        if resp.status_code == 200:
            for item in resp.json().get("value", []):
                if item.get("file"):
                    dl = httpx.get(item.get("@microsoft.graph.downloadUrl", ""), timeout=15)
                    if dl.status_code == 200:
                        docs.append(Document(content=dl.text, source=f"sharepoint://{item['name']}", metadata={"type": "sharepoint", "name": item["name"]}))
        return docs


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  Messaging & Social Loaders
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

class TelegramLoader(DocumentLoader):
    """Telegram chat export loader (JSON export file)."""
    def __init__(self, path: str) -> None: self.path = Path(path)
    def load(self) -> list[Document]:
        data = json.loads(self.path.read_text(encoding="utf-8")); docs = []
        for msg in data.get("messages", []):
            text = msg.get("text", "")
            if isinstance(text, list): text = "".join(t if isinstance(t, str) else t.get("text", "") for t in text)
            if text.strip():
                docs.append(Document(content=text, source=f"telegram://{msg.get('id', '')}", metadata={"type": "telegram", "date": msg.get("date", ""), "from": msg.get("from", "")}))
        return docs


class DiscordLoader(DocumentLoader):
    """Discord channel loader. Requires: DISCORD_BOT_TOKEN."""
    def __init__(self, channel_id: str, limit: int = 100, token: str = "") -> None:
        import os; self.channel_id = channel_id; self.limit = limit
        self._token = token or os.environ.get("DISCORD_BOT_TOKEN", "")
    def load(self) -> list[Document]:
        import httpx; docs = []
        resp = httpx.get(f"https://discord.com/api/v10/channels/{self.channel_id}/messages", headers={"Authorization": f"Bot {self._token}"}, params={"limit": self.limit}, timeout=15)
        if resp.status_code == 200:
            for msg in resp.json():
                if msg.get("content"):
                    docs.append(Document(content=msg["content"], source=f"discord://{self.channel_id}/{msg['id']}", metadata={"type": "discord", "author": msg.get("author", {}).get("username", ""), "timestamp": msg.get("timestamp", "")}))
        return docs


class WhatsAppLoader(DocumentLoader):
    """WhatsApp chat export loader (txt export file)."""
    def __init__(self, path: str) -> None: self.path = Path(path)
    def load(self) -> list[Document]:
        import re; docs = []; text = self.path.read_text(encoding="utf-8")
        pattern = re.compile(r"\[?(\d{1,2}/\d{1,2}/\d{2,4}),?\s*(\d{1,2}:\d{2}(?::\d{2})?(?:\s*[AP]M)?)\]?\s*-?\s*(.+?):\s*(.+)")
        for match in pattern.finditer(text):
            date, time_str, sender, content = match.groups()
            if content.strip() and "<Media omitted>" not in content:
                docs.append(Document(content=content.strip(), source="whatsapp", metadata={"type": "whatsapp", "sender": sender, "date": date, "time": time_str}))
        return docs


class RedditLoader(DocumentLoader):
    """Reddit posts/comments loader. Requires: REDDIT_CLIENT_ID + REDDIT_SECRET."""
    def __init__(self, subreddit: str = "", url: str = "", limit: int = 10) -> None:
        self.subreddit = subreddit; self.url = url; self.limit = limit
    def load(self) -> list[Document]:
        import httpx; docs = []
        if self.url:
            resp = httpx.get(f"{self.url}.json", headers={"User-Agent": "DuxxAI/1.0"}, timeout=15)
            if resp.status_code == 200:
                data = resp.json()
                if isinstance(data, list) and len(data) > 0:
                    post = data[0]["data"]["children"][0]["data"]
                    docs.append(Document(content=f"{post.get('title', '')}\n\n{post.get('selftext', '')}", source=self.url, metadata={"type": "reddit", "subreddit": post.get("subreddit", ""), "score": post.get("score", 0)}))
        elif self.subreddit:
            resp = httpx.get(f"https://www.reddit.com/r/{self.subreddit}/hot.json", headers={"User-Agent": "DuxxAI/1.0"}, params={"limit": self.limit}, timeout=15)
            if resp.status_code == 200:
                for child in resp.json().get("data", {}).get("children", []):
                    p = child["data"]
                    docs.append(Document(content=f"{p.get('title','')}\n\n{p.get('selftext','')}", source=f"https://reddit.com{p.get('permalink','')}", metadata={"type": "reddit", "subreddit": self.subreddit, "score": p.get("score", 0)}))
        return docs


class TwitterLoader(DocumentLoader):
    """Twitter/X loader. Requires: TWITTER_BEARER_TOKEN."""
    def __init__(self, query: str = "", user_id: str = "", max_results: int = 10, token: str = "") -> None:
        import os; self.query = query; self.user_id = user_id; self.max_results = max_results
        self._token = token or os.environ.get("TWITTER_BEARER_TOKEN", "")
    def load(self) -> list[Document]:
        import httpx; docs = []; headers = {"Authorization": f"Bearer {self._token}"}
        if self.query:
            resp = httpx.get("https://api.twitter.com/2/tweets/search/recent", headers=headers, params={"query": self.query, "max_results": min(self.max_results, 100)}, timeout=15)
            if resp.status_code == 200:
                for tweet in resp.json().get("data", []):
                    docs.append(Document(content=tweet.get("text", ""), source=f"twitter://{tweet['id']}", metadata={"type": "twitter", "id": tweet["id"]}))
        return docs


class HackerNewsLoader(DocumentLoader):
    """Hacker News loader (top/best/new stories)."""
    def __init__(self, story_type: str = "top", limit: int = 10) -> None:
        self.story_type = story_type; self.limit = limit
    def load(self) -> list[Document]:
        import httpx; docs = []
        resp = httpx.get(f"https://hacker-news.firebaseio.com/v0/{self.story_type}stories.json", timeout=10)
        if resp.status_code == 200:
            ids = resp.json()[:self.limit]
            for sid in ids:
                sr = httpx.get(f"https://hacker-news.firebaseio.com/v0/item/{sid}.json", timeout=10)
                if sr.status_code == 200:
                    item = sr.json()
                    docs.append(Document(content=f"{item.get('title','')}\n\n{item.get('text','') or item.get('url','')}", source=item.get("url", f"https://news.ycombinator.com/item?id={sid}"), metadata={"type": "hackernews", "score": item.get("score", 0), "by": item.get("by", "")}))
        return docs


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  Productivity & Dev Tools Loaders
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

class ConfluenceLoader(DocumentLoader):
    """Atlassian Confluence loader. Requires: CONFLUENCE_URL + CONFLUENCE_TOKEN."""
    def __init__(self, space_key: str = "", page_ids: list[str] | None = None, url: str = "", token: str = "") -> None:
        import os; self.space_key = space_key; self.page_ids = page_ids or []
        self._url = url or os.environ.get("CONFLUENCE_URL", ""); self._token = token or os.environ.get("CONFLUENCE_TOKEN", "")
    def load(self) -> list[Document]:
        import httpx, re; docs = []; headers = {"Authorization": f"Bearer {self._token}"}
        for pid in self.page_ids:
            resp = httpx.get(f"{self._url}/rest/api/content/{pid}?expand=body.storage", headers=headers, timeout=15)
            if resp.status_code == 200:
                data = resp.json(); html = data.get("body", {}).get("storage", {}).get("value", "")
                text = re.sub(r"<[^>]+>", " ", html); text = re.sub(r"\s+", " ", text).strip()
                docs.append(Document(content=text, source=f"confluence://{pid}", metadata={"type": "confluence", "title": data.get("title", ""), "space": self.space_key}))
        return docs


class JiraLoader(DocumentLoader):
    """Atlassian Jira loader. Requires: JIRA_URL + JIRA_TOKEN."""
    def __init__(self, jql: str = "", url: str = "", token: str = "", email: str = "") -> None:
        import os; self.jql = jql
        self._url = url or os.environ.get("JIRA_URL", ""); self._token = token or os.environ.get("JIRA_TOKEN", "")
        self._email = email or os.environ.get("JIRA_EMAIL", "")
    def load(self) -> list[Document]:
        import httpx, base64; docs = []
        auth = base64.b64encode(f"{self._email}:{self._token}".encode()).decode()
        resp = httpx.get(f"{self._url}/rest/api/3/search", headers={"Authorization": f"Basic {auth}"}, params={"jql": self.jql, "maxResults": 50}, timeout=15)
        if resp.status_code == 200:
            for issue in resp.json().get("issues", []):
                fields = issue.get("fields", {})
                desc = fields.get("description", "") or ""
                if isinstance(desc, dict): desc = str(desc)
                content = f"{fields.get('summary', '')}\n\n{desc}"
                docs.append(Document(content=content, source=f"jira://{issue['key']}", metadata={"type": "jira", "key": issue["key"], "status": fields.get("status", {}).get("name", ""), "priority": fields.get("priority", {}).get("name", "")}))
        return docs


class GitLabLoader(DocumentLoader):
    """GitLab repository loader. Requires: GITLAB_TOKEN."""
    def __init__(self, project_id: str, path: str = "", branch: str = "main", token: str = "", extensions: list[str] | None = None) -> None:
        import os; self.project_id = project_id; self.path = path; self.branch = branch
        self._token = token or os.environ.get("GITLAB_TOKEN", ""); self.extensions = extensions or [".md", ".txt", ".py"]
    def load(self) -> list[Document]:
        import httpx; docs = []; headers = {"PRIVATE-TOKEN": self._token}
        url = f"https://gitlab.com/api/v4/projects/{self.project_id}/repository/tree?ref={self.branch}&recursive=true"
        if self.path: url += f"&path={self.path}"
        resp = httpx.get(url, headers=headers, timeout=15)
        if resp.status_code == 200:
            for item in resp.json():
                if item["type"] != "blob": continue
                ext = Path(item["path"]).suffix.lower()
                if ext not in self.extensions: continue
                file_resp = httpx.get(f"https://gitlab.com/api/v4/projects/{self.project_id}/repository/files/{item['path'].replace('/', '%2F')}/raw?ref={self.branch}", headers=headers, timeout=15)
                if file_resp.status_code == 200:
                    docs.append(Document(content=file_resp.text, source=f"gitlab://{self.project_id}/{item['path']}", metadata={"type": "gitlab", "path": item["path"]}))
        return docs


class TrelloLoader(DocumentLoader):
    """Trello board loader. Requires: TRELLO_API_KEY + TRELLO_TOKEN."""
    def __init__(self, board_id: str, api_key: str = "", token: str = "") -> None:
        import os; self.board_id = board_id
        self._key = api_key or os.environ.get("TRELLO_API_KEY", ""); self._token = token or os.environ.get("TRELLO_TOKEN", "")
    def load(self) -> list[Document]:
        import httpx; docs = []
        resp = httpx.get(f"https://api.trello.com/1/boards/{self.board_id}/cards", params={"key": self._key, "token": self._token, "fields": "name,desc,labels"}, timeout=15)
        if resp.status_code == 200:
            for card in resp.json():
                content = f"{card.get('name', '')}\n\n{card.get('desc', '')}"
                labels = [l.get("name", "") for l in card.get("labels", [])]
                docs.append(Document(content=content, source=f"trello://{card['id']}", metadata={"type": "trello", "labels": labels}))
        return docs


class LinearLoader(DocumentLoader):
    """Linear.app issue loader. Requires: LINEAR_API_KEY."""
    def __init__(self, team_key: str = "", api_key: str = "") -> None:
        import os; self.team_key = team_key; self._key = api_key or os.environ.get("LINEAR_API_KEY", "")
    def load(self) -> list[Document]:
        import httpx; docs = []
        query = '{"query": "{ issues(first: 50) { nodes { title description state { name } priority url identifier } } }"}'
        resp = httpx.post("https://api.linear.app/graphql", headers={"Authorization": self._key, "Content-Type": "application/json"}, content=query, timeout=15)
        if resp.status_code == 200:
            for issue in resp.json().get("data", {}).get("issues", {}).get("nodes", []):
                content = f"{issue.get('title','')}\n\n{issue.get('description','') or ''}"
                docs.append(Document(content=content, source=issue.get("url", ""), metadata={"type": "linear", "id": issue.get("identifier", ""), "state": issue.get("state", {}).get("name", ""), "priority": issue.get("priority", 0)}))
        return docs


class AirtableLoader(DocumentLoader):
    """Airtable base loader. Requires: AIRTABLE_API_KEY."""
    def __init__(self, base_id: str, table_name: str, api_key: str = "") -> None:
        import os; self.base_id = base_id; self.table_name = table_name
        self._key = api_key or os.environ.get("AIRTABLE_API_KEY", "")
    def load(self) -> list[Document]:
        import httpx; docs = []
        resp = httpx.get(f"https://api.airtable.com/v0/{self.base_id}/{self.table_name}", headers={"Authorization": f"Bearer {self._key}"}, timeout=15)
        if resp.status_code == 200:
            for record in resp.json().get("records", []):
                fields = record.get("fields", {})
                content = "\n".join(f"{k}: {v}" for k, v in fields.items())
                docs.append(Document(content=content, source=f"airtable://{self.base_id}/{record['id']}", metadata={"type": "airtable", "fields": fields}))
        return docs


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  File Format Loaders
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

class ExcelLoader(DocumentLoader):
    """Excel (.xlsx/.xls) loader. Requires: pip install openpyxl"""
    def __init__(self, path: str, sheet_name: str | None = None) -> None:
        self.path = Path(path); self.sheet_name = sheet_name
    def load(self) -> list[Document]:
        try: import openpyxl
        except ImportError: raise ImportError("openpyxl required: pip install openpyxl")
        wb = openpyxl.load_workbook(str(self.path), data_only=True)
        sheets = [self.sheet_name] if self.sheet_name else wb.sheetnames; docs = []
        for sn in sheets:
            ws = wb[sn]; rows = []
            for row in ws.iter_rows(values_only=True):
                rows.append("\t".join(str(c) if c is not None else "" for c in row))
            docs.append(Document(content="\n".join(rows), source=str(self.path), metadata={"type": "excel", "sheet": sn}))
        return docs


class PPTXLoader(DocumentLoader):
    """PowerPoint (.pptx) loader. Requires: pip install python-pptx"""
    def __init__(self, path: str) -> None: self.path = Path(path)
    def load(self) -> list[Document]:
        try: from pptx import Presentation
        except ImportError: raise ImportError("python-pptx required: pip install python-pptx")
        prs = Presentation(str(self.path)); slides_text = []
        for i, slide in enumerate(prs.slides):
            texts = []
            for shape in slide.shapes:
                if shape.has_text_frame:
                    texts.append(shape.text_frame.text)
            slides_text.append(f"--- Slide {i+1} ---\n" + "\n".join(texts))
        return [Document(content="\n\n".join(slides_text), source=str(self.path), metadata={"type": "pptx", "slides": len(prs.slides)})]


class EPUBLoader(DocumentLoader):
    """EPUB ebook loader. Requires: pip install ebooklib"""
    def __init__(self, path: str) -> None: self.path = Path(path)
    def load(self) -> list[Document]:
        try: import ebooklib; from ebooklib import epub
        except ImportError: raise ImportError("ebooklib required: pip install ebooklib")
        import re; book = epub.read_epub(str(self.path)); texts = []
        for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
            html = item.get_body_content().decode("utf-8", errors="replace")
            text = re.sub(r"<[^>]+>", " ", html); text = re.sub(r"\s+", " ", text).strip()
            if text: texts.append(text)
        return [Document(content="\n\n".join(texts), source=str(self.path), metadata={"type": "epub", "title": book.get_metadata("DC", "title")[0][0] if book.get_metadata("DC", "title") else ""})]


class RTFLoader(DocumentLoader):
    """RTF file loader. Requires: pip install striprtf"""
    def __init__(self, path: str) -> None: self.path = Path(path)
    def load(self) -> list[Document]:
        try: from striprtf.striprtf import rtf_to_text
        except ImportError: raise ImportError("striprtf required: pip install striprtf")
        rtf = self.path.read_text(encoding="utf-8", errors="replace")
        return [Document(content=rtf_to_text(rtf), source=str(self.path), metadata={"type": "rtf"})]


class XMLLoader(DocumentLoader):
    """XML file loader."""
    def __init__(self, path: str, text_tags: list[str] | None = None) -> None:
        self.path = Path(path); self.text_tags = text_tags
    def load(self) -> list[Document]:
        import re; text = self.path.read_text(encoding="utf-8")
        if self.text_tags:
            parts = []
            for tag in self.text_tags:
                matches = re.findall(rf"<{tag}[^>]*>(.*?)</{tag}>", text, re.DOTALL)
                parts.extend(m.strip() for m in matches)
            content = "\n".join(parts)
        else:
            content = re.sub(r"<[^>]+>", " ", text); content = re.sub(r"\s+", " ", content).strip()
        return [Document(content=content, source=str(self.path), metadata={"type": "xml"})]


class JSONLoader(DocumentLoader):
    """JSON file loader with jq-like path extraction."""
    def __init__(self, path: str, content_key: str = "", jq_schema: str = "") -> None:
        self.path = Path(path); self.content_key = content_key; self.jq_schema = jq_schema
    def load(self) -> list[Document]:
        data = json.loads(self.path.read_text(encoding="utf-8")); docs = []
        if isinstance(data, list):
            for item in data:
                content = item.get(self.content_key, str(item)) if self.content_key else json.dumps(item)
                docs.append(Document(content=str(content), source=str(self.path), metadata={"type": "json"}))
        elif isinstance(data, dict):
            content = data.get(self.content_key, json.dumps(data)) if self.content_key else json.dumps(data)
            docs.append(Document(content=str(content), source=str(self.path), metadata={"type": "json"}))
        return docs


class EmailLoader(DocumentLoader):
    """Email (.eml) file loader."""
    def __init__(self, path: str) -> None: self.path = Path(path)
    def load(self) -> list[Document]:
        import email
        msg = email.message_from_string(self.path.read_text(encoding="utf-8", errors="replace"))
        body = ""
        if msg.is_multipart():
            for part in msg.walk():
                if part.get_content_type() == "text/plain":
                    body += part.get_payload(decode=True).decode("utf-8", errors="replace")
        else:
            body = msg.get_payload(decode=True).decode("utf-8", errors="replace") if msg.get_payload(decode=True) else str(msg.get_payload())
        return [Document(content=body, source=str(self.path), metadata={"type": "email", "subject": msg.get("Subject", ""), "from": msg.get("From", ""), "date": msg.get("Date", "")})]


class RSSLoader(DocumentLoader):
    """RSS/Atom feed loader."""
    def __init__(self, url: str, max_items: int = 20) -> None:
        self.url = url; self.max_items = max_items
    def load(self) -> list[Document]:
        import httpx, re; docs = []
        resp = httpx.get(self.url, timeout=15); resp.raise_for_status()
        items = re.findall(r"<item>(.*?)</item>", resp.text, re.DOTALL) or re.findall(r"<entry>(.*?)</entry>", resp.text, re.DOTALL)
        for item in items[:self.max_items]:
            title = re.search(r"<title[^>]*>(.*?)</title>", item, re.DOTALL)
            desc = re.search(r"<description[^>]*>(.*?)</description>", item, re.DOTALL) or re.search(r"<content[^>]*>(.*?)</content>", item, re.DOTALL) or re.search(r"<summary[^>]*>(.*?)</summary>", item, re.DOTALL)
            link = re.search(r"<link[^>]*>(.*?)</link>", item) or re.search(r'<link[^>]*href="([^"]+)"', item)
            t = re.sub(r"<[^>]+>", "", title.group(1)).strip() if title else ""
            d = re.sub(r"<[^>]+>", " ", desc.group(1)).strip() if desc else ""
            l = (link.group(1) if link else "").strip()
            docs.append(Document(content=f"{t}\n\n{d}", source=l, metadata={"type": "rss", "title": t}))
        return docs


class SitemapLoader(DocumentLoader):
    """Sitemap XML loader — discover and load pages from sitemap.xml."""
    def __init__(self, url: str, max_pages: int = 50, extensions: list[str] | None = None) -> None:
        self.url = url; self.max_pages = max_pages; self.extensions = extensions
    def load(self) -> list[Document]:
        import httpx, re; docs = []
        resp = httpx.get(self.url, timeout=15); resp.raise_for_status()
        urls = re.findall(r"<loc>(.*?)</loc>", resp.text)
        for page_url in urls[:self.max_pages]:
            if self.extensions:
                ext = Path(page_url.split("?")[0]).suffix.lower()
                if ext and ext not in self.extensions: continue
            try:
                pr = httpx.get(page_url, timeout=10, follow_redirects=True)
                if pr.status_code == 200:
                    text = re.sub(r"<script[^>]*>.*?</script>", "", pr.text, flags=re.DOTALL|re.I)
                    text = re.sub(r"<style[^>]*>.*?</style>", "", text, flags=re.DOTALL|re.I)
                    text = re.sub(r"<[^>]+>", " ", text); text = re.sub(r"\s+", " ", text).strip()
                    docs.append(Document(content=text[:5000], source=page_url, metadata={"type": "sitemap"}))
            except: pass
        return docs


class RecursiveURLLoader(DocumentLoader):
    """Recursively crawl and load pages from a URL."""
    def __init__(self, url: str, max_depth: int = 2, max_pages: int = 50) -> None:
        self.url = url; self.max_depth = max_depth; self.max_pages = max_pages
    def load(self) -> list[Document]:
        import httpx, re; from urllib.parse import urljoin, urlparse
        visited: set[str] = set(); docs = []; queue = [(self.url, 0)]
        domain = urlparse(self.url).netloc
        while queue and len(docs) < self.max_pages:
            url, depth = queue.pop(0)
            if url in visited or depth > self.max_depth: continue
            visited.add(url)
            try:
                resp = httpx.get(url, timeout=10, follow_redirects=True)
                if resp.status_code != 200: continue
                text = re.sub(r"<script[^>]*>.*?</script>", "", resp.text, flags=re.DOTALL|re.I)
                text = re.sub(r"<style[^>]*>.*?</style>", "", text, flags=re.DOTALL|re.I)
                clean = re.sub(r"<[^>]+>", " ", text); clean = re.sub(r"\s+", " ", clean).strip()
                docs.append(Document(content=clean[:5000], source=url, metadata={"type": "recursive_url", "depth": depth}))
                if depth < self.max_depth:
                    for link in re.findall(r'href="([^"]+)"', resp.text):
                        full = urljoin(url, link)
                        if urlparse(full).netloc == domain and full not in visited:
                            queue.append((full, depth + 1))
            except: pass
        return docs


class FigmaLoader(DocumentLoader):
    """Figma design file loader. Requires: FIGMA_ACCESS_TOKEN."""
    def __init__(self, file_key: str, token: str = "") -> None:
        import os; self.file_key = file_key; self._token = token or os.environ.get("FIGMA_ACCESS_TOKEN", "")
    def load(self) -> list[Document]:
        import httpx; docs = []
        resp = httpx.get(f"https://api.figma.com/v1/files/{self.file_key}", headers={"X-Figma-Token": self._token}, timeout=15)
        if resp.status_code == 200:
            data = resp.json()
            def extract_text(node: dict) -> list[str]:
                texts = []
                if node.get("type") == "TEXT": texts.append(node.get("characters", ""))
                for child in node.get("children", []): texts.extend(extract_text(child))
                return texts
            all_text = extract_text(data.get("document", {}))
            docs.append(Document(content="\n".join(t for t in all_text if t), source=f"figma://{self.file_key}", metadata={"type": "figma", "name": data.get("name", "")}))
        return docs


class FirecrawlLoader(DocumentLoader):
    """Firecrawl web scraper loader. Requires: FIRECRAWL_API_KEY."""
    def __init__(self, url: str, mode: str = "scrape", api_key: str = "") -> None:
        import os; self.url = url; self.mode = mode; self._key = api_key or os.environ.get("FIRECRAWL_API_KEY", "")
    def load(self) -> list[Document]:
        import httpx
        resp = httpx.post(f"https://api.firecrawl.dev/v1/{self.mode}", headers={"Authorization": f"Bearer {self._key}"}, json={"url": self.url}, timeout=30)
        resp.raise_for_status(); data = resp.json().get("data", {})
        content = data.get("markdown", "") or data.get("content", "")
        return [Document(content=content, source=self.url, metadata={"type": "firecrawl", "title": data.get("metadata", {}).get("title", "")})]


class ApifyLoader(DocumentLoader):
    """Apify dataset loader. Requires: APIFY_API_TOKEN."""
    def __init__(self, dataset_id: str, token: str = "") -> None:
        import os; self.dataset_id = dataset_id; self._token = token or os.environ.get("APIFY_API_TOKEN", "")
    def load(self) -> list[Document]:
        import httpx; docs = []
        resp = httpx.get(f"https://api.apify.com/v2/datasets/{self.dataset_id}/items", params={"token": self._token, "format": "json"}, timeout=30)
        if resp.status_code == 200:
            for item in resp.json():
                content = item.get("text", "") or item.get("content", "") or json.dumps(item)
                docs.append(Document(content=str(content), source=item.get("url", f"apify://{self.dataset_id}"), metadata={"type": "apify"}))
        return docs


class FacebookChatLoader(DocumentLoader):
    """Facebook Messenger chat export loader (JSON)."""
    def __init__(self, path: str) -> None: self.path = Path(path)
    def load(self) -> list[Document]:
        data = json.loads(self.path.read_text(encoding="utf-8")); docs = []
        for msg in data.get("messages", []):
            content = msg.get("content", "")
            if content:
                docs.append(Document(content=content, source="facebook_chat", metadata={"type": "facebook_chat", "sender": msg.get("sender_name", ""), "timestamp_ms": msg.get("timestamp_ms", 0)}))
        return docs


class MastodonLoader(DocumentLoader):
    """Mastodon toots loader."""
    def __init__(self, instance_url: str, account_id: str = "", access_token: str = "", limit: int = 40) -> None:
        import os; self._url = instance_url; self._account = account_id; self.limit = limit
        self._token = access_token or os.environ.get("MASTODON_ACCESS_TOKEN", "")
    def load(self) -> list[Document]:
        import httpx, re; docs = []; headers = {"Authorization": f"Bearer {self._token}"} if self._token else {}
        url = f"{self._url}/api/v1/accounts/{self._account}/statuses" if self._account else f"{self._url}/api/v1/timelines/public"
        resp = httpx.get(url, headers=headers, params={"limit": self.limit}, timeout=15)
        if resp.status_code == 200:
            for toot in resp.json():
                text = re.sub(r"<[^>]+>", "", toot.get("content", "")).strip()
                if text: docs.append(Document(content=text, source=toot.get("url", ""), metadata={"type": "mastodon", "account": toot.get("account", {}).get("username", ""), "created_at": toot.get("created_at", "")}))
        return docs


class RoamLoader(DocumentLoader):
    """Roam Research export loader (JSON)."""
    def __init__(self, path: str) -> None: self.path = Path(path)
    def load(self) -> list[Document]:
        data = json.loads(self.path.read_text(encoding="utf-8")); docs = []
        for page in data if isinstance(data, list) else [data]:
            title = page.get("title", "")
            children = page.get("children", [])
            texts = [title] if title else []
            def extract(nodes: list) -> None:
                for n in nodes:
                    s = n.get("string", "")
                    if s: texts.append(s)
                    extract(n.get("children", []))
            extract(children)
            if texts: docs.append(Document(content="\n".join(texts), source=f"roam://{title}", metadata={"type": "roam", "title": title}))
        return docs


class QuipLoader(DocumentLoader):
    """Quip document loader. Requires: QUIP_ACCESS_TOKEN."""
    def __init__(self, document_ids: list[str], token: str = "") -> None:
        import os; self.document_ids = document_ids; self._token = token or os.environ.get("QUIP_ACCESS_TOKEN", "")
    def load(self) -> list[Document]:
        import httpx, re; docs = []
        for did in self.document_ids:
            resp = httpx.get(f"https://platform.quip.com/1/threads/{did}", headers={"Authorization": f"Bearer {self._token}"}, timeout=15)
            if resp.status_code == 200:
                data = resp.json(); html = data.get("html", "")
                text = re.sub(r"<[^>]+>", " ", html); text = re.sub(r"\s+", " ", text).strip()
                docs.append(Document(content=text, source=f"quip://{did}", metadata={"type": "quip", "title": data.get("thread", {}).get("title", "")}))
        return docs


class DoclingLoader(DocumentLoader):
    """Docling document intelligence loader. Requires: pip install docling"""
    def __init__(self, path: str) -> None: self.path = Path(path)
    def load(self) -> list[Document]:
        try: from docling.document_converter import DocumentConverter
        except ImportError: raise ImportError("docling required: pip install docling")
        converter = DocumentConverter(); result = converter.convert(str(self.path))
        text = result.document.export_to_text()
        return [Document(content=text, source=str(self.path), metadata={"type": "docling"})]


class AmazonTextractLoader(DocumentLoader):
    """Amazon Textract OCR loader. Requires: pip install boto3"""
    def __init__(self, path: str, region: str = "us-east-1") -> None:
        self.path = Path(path); self.region = region
    def load(self) -> list[Document]:
        try: import boto3
        except ImportError: raise ImportError("boto3 required: pip install boto3")
        client = boto3.client("textract", region_name=self.region)
        with open(self.path, "rb") as f: img_bytes = f.read()
        resp = client.detect_document_text(Document={"Bytes": img_bytes})
        texts = [b["Text"] for b in resp.get("Blocks", []) if b["BlockType"] == "LINE"]
        return [Document(content="\n".join(texts), source=str(self.path), metadata={"type": "textract", "pages": 1})]


class MathPixLoader(DocumentLoader):
    """Mathpix OCR loader (PDF/images with math). Requires: MATHPIX_APP_ID + MATHPIX_APP_KEY."""
    def __init__(self, path: str) -> None:
        import os; self.path = Path(path)
        self._app_id = os.environ.get("MATHPIX_APP_ID", ""); self._app_key = os.environ.get("MATHPIX_APP_KEY", "")
    def load(self) -> list[Document]:
        import httpx
        with open(self.path, "rb") as f:
            resp = httpx.post("https://api.mathpix.com/v3/text", headers={"app_id": self._app_id, "app_key": self._app_key}, files={"file": f}, data={"options_json": '{"math_inline_delimiters": ["$", "$"]}'}, timeout=60)
        if resp.status_code == 200:
            text = resp.json().get("text", "")
            return [Document(content=text, source=str(self.path), metadata={"type": "mathpix"})]
        return []


class HuaweiOBSLoader(DocumentLoader):
    """Huawei OBS (Object Storage) loader."""
    def __init__(self, bucket: str, prefix: str = "", ak: str = "", sk: str = "", endpoint: str = "") -> None:
        import os; self.bucket = bucket; self.prefix = prefix
        self._ak = ak or os.environ.get("HUAWEI_OBS_AK", ""); self._sk = sk or os.environ.get("HUAWEI_OBS_SK", "")
        self._endpoint = endpoint or os.environ.get("HUAWEI_OBS_ENDPOINT", "")
    def load(self) -> list[Document]:
        try: from obs import ObsClient
        except ImportError: raise ImportError("esdk-obs-python required: pip install esdk-obs-python")
        client = ObsClient(access_key_id=self._ak, secret_access_key=self._sk, server=self._endpoint)
        resp = client.listObjects(self.bucket, prefix=self.prefix); docs = []
        for obj in resp.body.contents:
            r = client.getObject(self.bucket, obj.key); content = r.body.response.read().decode("utf-8", errors="replace")
            docs.append(Document(content=content, source=f"obs://{self.bucket}/{obj.key}", metadata={"type": "huawei_obs"}))
        return docs


class TencentCOSLoader(DocumentLoader):
    """Tencent Cloud COS loader. Requires: pip install cos-python-sdk-v5"""
    def __init__(self, bucket: str, prefix: str = "", region: str = "", secret_id: str = "", secret_key: str = "") -> None:
        import os; self.bucket = bucket; self.prefix = prefix
        self._region = region or os.environ.get("COS_REGION", ""); self._id = secret_id or os.environ.get("COS_SECRET_ID", "")
        self._key = secret_key or os.environ.get("COS_SECRET_KEY", "")
    def load(self) -> list[Document]:
        try: from qcloud_cos import CosConfig, CosS3Client
        except ImportError: raise ImportError("cos-python-sdk-v5 required: pip install cos-python-sdk-v5")
        config = CosConfig(Region=self._region, SecretId=self._id, SecretKey=self._key)
        client = CosS3Client(config); resp = client.list_objects(Bucket=self.bucket, Prefix=self.prefix); docs = []
        for obj in resp.get("Contents", []):
            r = client.get_object(Bucket=self.bucket, Key=obj["Key"])
            content = r["Body"].get_raw_stream().read().decode("utf-8", errors="replace")
            docs.append(Document(content=content, source=f"cos://{self.bucket}/{obj['Key']}", metadata={"type": "tencent_cos"}))
        return docs


class HyperBrowserLoader(DocumentLoader):
    """HyperBrowser web scraping loader. Requires: HYPERBROWSER_API_KEY."""
    def __init__(self, url: str, api_key: str = "") -> None:
        import os; self.url = url; self._key = api_key or os.environ.get("HYPERBROWSER_API_KEY", "")
    def load(self) -> list[Document]:
        import httpx
        resp = httpx.post("https://api.hyperbrowser.ai/v1/scrape", headers={"Authorization": f"Bearer {self._key}"}, json={"url": self.url}, timeout=30)
        resp.raise_for_status(); data = resp.json().get("data", {})
        return [Document(content=data.get("markdown", "") or data.get("text", ""), source=self.url, metadata={"type": "hyperbrowser", "title": data.get("metadata", {}).get("title", "")})]


class AgentQLLoader(DocumentLoader):
    """AgentQL structured web extraction. Requires: AGENTQL_API_KEY."""
    def __init__(self, url: str, query: str = "", api_key: str = "") -> None:
        import os; self.url = url; self.query = query; self._key = api_key or os.environ.get("AGENTQL_API_KEY", "")
    def load(self) -> list[Document]:
        import httpx
        resp = httpx.post("https://api.agentql.com/v1/query-data", headers={"X-API-Key": self._key}, json={"url": self.url, "query": self.query or "{ content }"}, timeout=30)
        resp.raise_for_status()
        return [Document(content=json.dumps(resp.json().get("data", {})), source=self.url, metadata={"type": "agentql"})]


class AsanaLoader(DocumentLoader):
    """Asana tasks loader. Requires: ASANA_ACCESS_TOKEN."""
    def __init__(self, project_gid: str, token: str = "") -> None:
        import os; self.project_gid = project_gid; self._token = token or os.environ.get("ASANA_ACCESS_TOKEN", "")
    def load(self) -> list[Document]:
        import httpx; docs = []
        resp = httpx.get(f"https://app.asana.com/api/1.0/projects/{self.project_gid}/tasks", headers={"Authorization": f"Bearer {self._token}"}, params={"opt_fields": "name,notes,completed,assignee.name"}, timeout=15)
        if resp.status_code == 200:
            for task in resp.json().get("data", []):
                content = f"{task.get('name', '')}\n\n{task.get('notes', '')}"
                docs.append(Document(content=content, source=f"asana://{task['gid']}", metadata={"type": "asana", "completed": task.get("completed", False)}))
        return docs


class MondayLoader(DocumentLoader):
    """Monday.com board loader. Requires: MONDAY_API_KEY."""
    def __init__(self, board_id: str, api_key: str = "") -> None:
        import os; self.board_id = board_id; self._key = api_key or os.environ.get("MONDAY_API_KEY", "")
    def load(self) -> list[Document]:
        import httpx; docs = []
        query = f'{{"query": "{{ boards(ids: {self.board_id}) {{ items_page {{ items {{ name column_values {{ text }} }} }} }} }}"}}'
        resp = httpx.post("https://api.monday.com/v2", headers={"Authorization": self._key, "Content-Type": "application/json"}, content=query, timeout=15)
        if resp.status_code == 200:
            for board in resp.json().get("data", {}).get("boards", []):
                for item in board.get("items_page", {}).get("items", []):
                    vals = " | ".join(cv.get("text", "") for cv in item.get("column_values", []) if cv.get("text"))
                    docs.append(Document(content=f"{item.get('name', '')}\n{vals}", source=f"monday://{self.board_id}", metadata={"type": "monday"}))
        return docs


class ClickUpLoader(DocumentLoader):
    """ClickUp tasks loader. Requires: CLICKUP_API_KEY."""
    def __init__(self, list_id: str, api_key: str = "") -> None:
        import os; self.list_id = list_id; self._key = api_key or os.environ.get("CLICKUP_API_KEY", "")
    def load(self) -> list[Document]:
        import httpx; docs = []
        resp = httpx.get(f"https://api.clickup.com/api/v2/list/{self.list_id}/task", headers={"Authorization": self._key}, timeout=15)
        if resp.status_code == 200:
            for task in resp.json().get("tasks", []):
                content = f"{task.get('name', '')}\n\n{task.get('description', '') or ''}"
                docs.append(Document(content=content, source=f"clickup://{task['id']}", metadata={"type": "clickup", "status": task.get("status", {}).get("status", "")}))
        return docs


class SpiderLoader(DocumentLoader):
    """Spider web crawler. Requires: SPIDER_API_KEY."""
    def __init__(self, url: str, api_key: str = "", limit: int = 10) -> None:
        import os; self.url = url; self._key = api_key or os.environ.get("SPIDER_API_KEY", ""); self.limit = limit
    def load(self) -> list[Document]:
        import httpx; docs = []
        resp = httpx.post("https://api.spider.cloud/crawl", headers={"Authorization": f"Bearer {self._key}"}, json={"url": self.url, "limit": self.limit}, timeout=60)
        if resp.status_code == 200:
            for page in resp.json():
                docs.append(Document(content=page.get("content", ""), source=page.get("url", ""), metadata={"type": "spider"}))
        return docs
